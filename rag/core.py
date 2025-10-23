import os
import re
import json
from pathlib import Path
from typing import Dict, List, Optional, Callable

# --- NUEVAS DEPENDENCIAS (CON CORRECCIÓN DE PINECONE) ---
from pinecone import Pinecone, ServerlessSpec
from langchain_pinecone import PineconeVectorStore
from langchain_google_genai import GoogleGenerativeAIEmbeddings, ChatGoogleGenerativeAI
from langchain_core.documents import Document as LCDocument
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.runnables import RunnablePassthrough
from langchain_core.output_parsers import StrOutputParser

# --- DEPENDENCIAS EXISTENTES ---
from pypdf import PdfReader
from docx import Document as DocxDocument

# --- PARSERS DE ARCHIVOS (sin cambios) ---
def _read_txt(p: Path) -> str: return p.read_text(encoding="utf-8", errors="ignore")
def _read_pdf(p: Path) -> str:
    r = PdfReader(str(p))
    return "\n".join(page.extract_text() or "" for page in r.pages)
def _read_docx(p: Path) -> str:
    d = DocxDocument(str(p))
    return "\n".join(p.text for p in d.paragraphs)

READERS = { ".txt": _read_txt, ".md": _read_txt, ".pdf": _read_pdf, ".docx": _read_docx}

# --- LÓGICA DE METADATA (sin cambios) ---
def get_document_metadata(file_path: Path) -> Dict:
    """Crea metadata rica para un documento, incluyendo tipo y año."""
    suffix = file_path.suffix.lower()
    metadata = {
        "file_path": str(file_path),
        "file_name": file_path.name,
        "file_extension": suffix,
        "source_type": "unknown"
    }
    if suffix in [ ".pdf", ".docx", ".pptx", ".xlsx"]:
        metadata["source_type"] = "document"
    elif suffix == ".txt":
        if 'transcripciones_tiktok' in str(file_path).lower():
            metadata["source_type"] = "video_audio"
        else:
            metadata["source_type"] = "txt_md"
    year_match = re.search(r'(202[0-9])', str(file_path))
    if year_match:
        metadata["year"] = int(year_match.group(1))
    
    if "transcripciones_tiktok" in str(file_path):
        print(f"DEBUG: Metadata for {file_path}: {metadata}")

    return metadata

# --- CLASE PRINCIPAL DEL SISTEMA RAG (CORREGIDA) ---
class RAGSystem:
    def __init__(self,
                 embeddings_model_name: str = "models/embedding-001",
                 llm_model_name: str = "models/gemini-2.5-pro",
                 pinecone_index_name: str = "chatbot-rag-gemini",
                 data_dir: str = "./data",
                 progress_cb: Optional[Callable] = None):

        self.llm_model_name = llm_model_name
        self.pinecone_index_name = pinecone_index_name
        self.top_k = 5
        self.data_dir = Path(data_dir)
        self._progress = progress_cb or (lambda *_, **__: None)

        self.embed_fn = GoogleGenerativeAIEmbeddings(model=embeddings_model_name, task_type="retrieval_query")
        self.vector_store = self._init_pinecone()
        self.llm = ChatGoogleGenerativeAI(model=self.llm_model_name, temperature=0.1, convert_system_message_to_human=True)
        self.rag_chain = self._create_rag_chain()

    def _init_pinecone(self) -> PineconeVectorStore:
        """Inicializa la conexión a Pinecone usando la nueva sintaxis."""
        pc = Pinecone(api_key=os.getenv("PINECONE_API_KEY"))

        if self.pinecone_index_name not in pc.list_indexes().names():
            print(f"El índice '{self.pinecone_index_name}' no existe. Construyendo desde cero...")
            self._build_and_upload_index(pinecone_client=pc)
        else:
            print(f"Índice '{self.pinecone_index_name}' encontrado. Cargando...")

        return PineconeVectorStore.from_existing_index(self.pinecone_index_name, self.embed_fn)

    def _chunk_text(self, text: str, size: int = 1000, overlap: int = 150) -> List[str]:
        return [text[i:i+size] for i in range(0, len(text), size - overlap)]

    def _build_and_upload_index(self, pinecone_client: Pinecone):
        """Crea un índice en Pinecone y sube los documentos."""
        all_chunks = []
        files_to_process = [p for p in self.data_dir.rglob("*") if p.suffix.lower() in READERS and p.is_file()]
        total_files = len(files_to_process)

        for i, file_path in enumerate(files_to_process):
            self._progress("Procesando archivos", i + 1, total_files)
            try:
                content = READERS[file_path.suffix.lower()](file_path)
                if not content: continue
                metadata = get_document_metadata(file_path)
                chunks = self._chunk_text(content)
                for chunk_content in chunks:
                    all_chunks.append(LCDocument(page_content=chunk_content, metadata=metadata))
            except Exception as e:
                print(f"Error procesando {file_path}: {e}")

        if not all_chunks:
            print("No se encontraron documentos para indexar.")
            return

        print(f"Creando índice '{self.pinecone_index_name}' en Pinecone...")
        # La dimensión para 'models/embedding-001' de Gemini es 768
        embedding_dimension = 768
        pinecone_client.create_index(
            name=self.pinecone_index_name,
            dimension=embedding_dimension,
            metric='cosine',
            spec=ServerlessSpec(cloud='aws', region='us-east-1') # Spec recomendada
        )

        print(f"Subiendo {len(all_chunks)} chunks al índice...")
        PineconeVectorStore.from_documents(
            documents=all_chunks,
            embedding=self.embed_fn,
            index_name=self.pinecone_index_name
        )
        print("Índice creado y documentos subidos.")

    def rebuild_index(self):
        """Deletes the existing Pinecone index and rebuilds it."""
        pc = Pinecone(api_key=os.getenv("PINECONE_API_KEY"))
        if self.pinecone_index_name in pc.list_indexes().names():
            print(f"Deleting existing index '{self.pinecone_index_name}'...")
            pc.delete_index(self.pinecone_index_name)
            print("Index deleted.")
        self._build_and_upload_index(pinecone_client=pc)

    def _create_rag_chain(self):
        template = """
Eres un asistente experto de la agencia Labelium. Tu nombre es Labelix.
Responde a la pregunta del usuario basándote ESTRICTA Y ÚNICAMENTE en el siguiente contexto.
Al final de tu respuesta, cita TODAS las fuentes que has usado de la metadata en una sección llamada 'Fuentes:'.
Si un documento en el contexto tiene `"source_type": "video_audio"`, DEBES indicar que la información proviene de una transcripción de video.

CONTEXTO:
{context}

PREGUNTA:
{question}

RESPUESTA:
"""
        prompt_template = ChatPromptTemplate.from_template(template)
        retriever = self.vector_store.as_retriever(search_kwargs={'k': self.top_k})

        def format_docs(docs: List[LCDocument]) -> str:
            formatted_docs = []
            for doc in docs:
                metadata_str = json.dumps(doc.metadata, ensure_ascii=False, indent=2)
                formatted_docs.append(f"Contenido: {doc.page_content}\nMetadata:\n{metadata_str}")
            print("DEBUG: Formatted context for LLM:")
            print("\n\n---\n\n".join(formatted_docs))
            
            return "\n\n---\n\n".join(formatted_docs)

        rag_chain = (
            {"context": retriever | format_docs, "question": RunnablePassthrough()}
            | prompt_template
            | self.llm
            | StrOutputParser()
        )
        return rag_chain

    def query(self, prompt: str) -> Dict:
        if not prompt:
            return {"answer": "Por favor, haz una pregunta.", "sources": []}

        # La cadena ahora se invoca con el prompt (string) directamente.
        answer = self.rag_chain.invoke(prompt)
        
        # Para encontrar las fuentes, necesitamos ejecutar el retriever por separado.
        retrieved_docs = self.vector_store.as_retriever(search_kwargs={'k': self.top_k}).invoke(prompt)
        sources = list(set(doc.metadata.get("file_path", "") for doc in retrieved_docs if doc.metadata))

        return {"answer": answer, "sources": sources}
